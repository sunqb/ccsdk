"""Document parser for RAG ingestion."""
from __future__ import annotations

import logging
from dataclasses import dataclass, field
from io import BytesIO
from pathlib import Path
from typing import Any, Protocol, runtime_checkable

import httpx

SUPPORTED_TEXT_EXTENSIONS = {".txt", ".md"}
SUPPORTED_MINERU_EXTENSIONS = {".pdf", ".docx"}
SUPPORTED_ALL_EXTENSIONS = SUPPORTED_TEXT_EXTENSIONS | SUPPORTED_MINERU_EXTENSIONS

MIME_TYPES = {
    ".txt": "text/plain",
    ".md": "text/markdown",
    ".pdf": "application/pdf",
    ".docx": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
}

log = logging.getLogger(__name__)

_NETWORK_RETRY_EXCEPTIONS = (
    httpx.ConnectError,
    httpx.ConnectTimeout,
    httpx.NetworkError,
    httpx.ReadTimeout,
    httpx.RemoteProtocolError,
)


@dataclass(slots=True)
class ParsedDocument:
    """Parsed text document with lightweight metadata."""

    filename: str
    mime_type: str
    text: str
    metadata: dict[str, Any] = field(default_factory=dict)


@runtime_checkable
class DocumentParser(Protocol):
    """Pluggable document parser protocol.

    Implement this protocol to add new parsing backends (e.g. MinerU, LlamaParse).
    """

    supported_extensions: set[str]

    def parse_bytes(
        self,
        content: bytes,
        *,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> ParsedDocument:
        """Parse a supported document from bytes."""
        ...


class LocalDocumentParser:
    """Parse supported documents locally into plain text for RAG indexing.

    This is the default local parser used when RAG_PARSER_PROVIDER=local. It
    handles .txt/.md directly and keeps PDF/DOCX support as a local fallback for
    development or deployments that explicitly choose the local provider.
    """

    supported_extensions = SUPPORTED_ALL_EXTENSIONS

    def parse_bytes(
        self,
        content: bytes,
        *,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> ParsedDocument:
        """Parse a supported document from bytes using local dependencies."""
        suffix = Path(filename).suffix.lower()
        self._ensure_supported(suffix)
        if suffix in SUPPORTED_TEXT_EXTENSIONS:
            text = self._decode_text(content)
        elif suffix == ".pdf":
            text = MinerUDocumentParser._parse_pdf_local(content)
        elif suffix == ".docx":
            text = MinerUDocumentParser._parse_docx_local(content)
        else:
            supported = ", ".join(sorted(self.supported_extensions))
            raise ValueError(f"Unsupported local document type: {suffix}. Supported: {supported}")
        return ParsedDocument(
            filename=filename,
            mime_type=MIME_TYPES.get(suffix, "text/plain"),
            text=self._normalize_newlines(text),
            metadata={**(metadata or {}), "extension": suffix, "parser": "local"},
        )

    def _ensure_supported(self, suffix: str) -> None:
        if suffix not in self.supported_extensions:
            supported = ", ".join(sorted(self.supported_extensions))
            raise ValueError(f"Unsupported local document type: {suffix}. Supported: {supported}")

    @staticmethod
    def _decode_text(content: bytes) -> str:
        try:
            return content.decode("utf-8-sig")
        except UnicodeDecodeError:
            return content.decode("utf-8", errors="replace")

    @staticmethod
    def _normalize_newlines(text: str) -> str:
        return text.replace("\r\n", "\n").replace("\r", "\n")


class MinerUDocumentParser:
    """Parse .pdf / .docx documents via the MinerU HTTP service.

    MinerU returns structured Markdown which preserves heading hierarchy,
    enabling the TextChunker to generate accurate parent-child metadata.
    """

    supported_extensions = SUPPORTED_MINERU_EXTENSIONS

    def __init__(
        self,
        *,
        base_url: str,
        api_key: str | None = None,
        timeout_seconds: float = 120.0,
        fallback_to_local: bool = False,
        client: httpx.Client | None = None,
    ) -> None:
        self.base_url = base_url.rstrip("/")
        self.api_key = api_key or ""
        self.timeout_seconds = timeout_seconds
        self.fallback_to_local = fallback_to_local
        self._client = client

    def parse_bytes(
        self,
        content: bytes,
        *,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> ParsedDocument:
        """Parse a .pdf or .docx document via MinerU HTTP API."""
        suffix = Path(filename).suffix.lower()
        self._ensure_supported(suffix)

        try:
            text, resp_metadata = self._call_mineru(content, filename, suffix)
        except Exception as exc:
            if self.fallback_to_local:
                log.warning("MinerU failed for %s, falling back to local parser: %s", filename, exc)
                return self._local_fallback(content, filename=filename, metadata=metadata)
            raise ValueError(f"MinerU parsing failed for {filename}: {exc}") from exc

        return ParsedDocument(
            filename=filename,
            mime_type=MIME_TYPES.get(suffix, "application/octet-stream"),
            text=text,
            metadata={
                **(metadata or {}),
                "extension": suffix,
                "parser": "mineru",
                **resp_metadata,
            },
        )

    def _call_mineru(
        self,
        content: bytes,
        filename: str,
        suffix: str,
    ) -> tuple[str, dict[str, Any]]:
        headers = {"Authorization": f"Bearer {self.api_key}"} if self.api_key else {}
        files = {"file": (filename, content, MIME_TYPES.get(suffix, "application/octet-stream"))}
        data = {"to_markdown": "true"}

        last_exc: Exception | None = None
        for attempt in range(2):
            try:
                resp = self._post_mineru(headers=headers, files=files, data=data)
                break
            except _NETWORK_RETRY_EXCEPTIONS as exc:
                last_exc = exc
                if attempt == 0:
                    log.warning("Retrying MinerU parse for %s after network error: %s", filename, exc)
                    continue
                raise ValueError("MinerU network error") from exc
        else:
            raise ValueError("MinerU network error") from last_exc

        if resp.status_code >= 400:
            log.warning(
                "MinerU HTTP error for %s: status=%s body=%s",
                filename,
                resp.status_code,
                resp.text[:500],
            )
            raise ValueError(f"MinerU returned HTTP {resp.status_code}")

        result: Any
        try:
            result = resp.json()
        except ValueError:
            result = resp.text

        # MinerU may return Markdown directly or wrapped in a result field
        if isinstance(result, dict):
            text = result.get("markdown", result.get("text", result.get("content", "")))
            page_count = result.get("page_count", result.get("pageCount"))
            document_id = result.get("document_id", result.get("documentId"))
            parser_version = result.get("version", result.get("mineru_version"))
        else:
            text = str(result)
            page_count = None
            document_id = None
            parser_version = None

        resp_metadata: dict[str, Any] = {}
        if parser_version is not None:
            resp_metadata["parserVersion"] = parser_version
        if page_count is not None:
            resp_metadata["pageCount"] = page_count
        if document_id is not None:
            resp_metadata["documentId"] = document_id

        if not text or not text.strip():
            raise ValueError(f"MinerU returned empty content for {filename}")

        return text, resp_metadata

    def _post_mineru(
        self,
        *,
        headers: dict[str, str],
        files: dict[str, tuple[str, bytes, str]],
        data: dict[str, str],
    ) -> httpx.Response:
        if self._client is not None:
            return self._client.post(
                f"{self.base_url}/parse",
                headers=headers,
                files=files,
                data=data,
            )

        with httpx.Client(timeout=self.timeout_seconds) as client:
            return client.post(
                f"{self.base_url}/parse",
                headers=headers,
                files=files,
                data=data,
            )

    def _local_fallback(
        self,
        content: bytes,
        *,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        if suffix == ".pdf":
            text = self._parse_pdf_local(content)
        elif suffix == ".docx":
            text = self._parse_docx_local(content)
        else:
            raise ValueError(f"No local fallback available for {suffix}")

        return ParsedDocument(
            filename=filename,
            mime_type=MIME_TYPES.get(suffix, "text/plain"),
            text=self._normalize_newlines(text),
            metadata={
                **(metadata or {}),
                "extension": suffix,
                "parser": "local",
                "parserFallbackFrom": "mineru",
            },
        )

    @staticmethod
    def _parse_pdf_local(content: bytes) -> str:
        try:
            from pypdf import PdfReader
        except ImportError as exc:
            raise ValueError("PDF parsing requires optional dependency 'pypdf'") from exc

        reader = PdfReader(BytesIO(content))
        pages = [page.extract_text() or "" for page in reader.pages]
        text = "\n\n".join(page.strip() for page in pages if page.strip())
        if not text:
            raise ValueError("No text content found in PDF document")
        return text

    @staticmethod
    def _parse_docx_local(content: bytes) -> str:
        try:
            from docx import Document
        except ImportError as exc:
            raise ValueError("DOCX parsing requires optional dependency 'python-docx'") from exc

        document = Document(BytesIO(content))
        blocks = [paragraph.text.strip() for paragraph in document.paragraphs]
        for table in document.tables:
            for row in table.rows:
                cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if cells:
                    blocks.append(" | ".join(cells))

        text = "\n\n".join(block for block in blocks if block)
        if not text:
            raise ValueError("No text content found in DOCX document")
        return text

    def _ensure_supported(self, suffix: str) -> None:
        if suffix not in self.supported_extensions:
            supported = ", ".join(sorted(self.supported_extensions))
            raise ValueError(f"Unsupported MinerU document type: {suffix}. Supported: {supported}")

    @staticmethod
    def _normalize_newlines(text: str) -> str:
        return text.replace("\r\n", "\n").replace("\r", "\n")


class HybridDocumentParser:
    """Router that delegates to LocalDocumentParser or MinerUDocumentParser.

    Use this as the single parsing entry point when both local and MinerU
    backends are needed in the same service.
    """

    def __init__(
        self,
        *,
        mineru_base_url: str | None = None,
        mineru_api_key: str | None = None,
        mineru_timeout_seconds: float = 120.0,
        mineru_fallback_to_local: bool = False,
    ) -> None:
        self.local = LocalDocumentParser()
        self.mineru: MinerUDocumentParser | None = None

        if mineru_base_url:
            self.mineru = MinerUDocumentParser(
                base_url=mineru_base_url,
                api_key=mineru_api_key,
                timeout_seconds=mineru_timeout_seconds,
                fallback_to_local=mineru_fallback_to_local,
            )

    @property
    def supported_extensions(self) -> set[str]:
        exts = set(SUPPORTED_TEXT_EXTENSIONS)
        if self.mineru:
            exts |= self.mineru.supported_extensions
        return exts

    def parse_bytes(
        self,
        content: bytes,
        *,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> ParsedDocument:
        """Route to the appropriate parser based on file extension."""
        suffix = Path(filename).suffix.lower()

        if suffix in SUPPORTED_TEXT_EXTENSIONS:
            return self.local.parse_bytes(content, filename=filename, metadata=metadata)

        if suffix in SUPPORTED_MINERU_EXTENSIONS:
            if self.mineru is None:
                raise ValueError(
                    f"MinerU is not configured but is required for {suffix} files. "
                    "Set RAG_PARSER_PROVIDER=mineru and MINERU_BASE_URL."
                )
            return self.mineru.parse_bytes(content, filename=filename, metadata=metadata)

        supported = ", ".join(sorted(self.supported_extensions))
        raise ValueError(f"Unsupported document type: {suffix}. Supported: {supported}")


class TextDocumentParser(LocalDocumentParser):
    """Backward-compatible local parser for legacy tests and callers.

    New production wiring should use LocalDocumentParser for .txt/.md and
    MinerUDocumentParser/HybridDocumentParser for .pdf/.docx. This class keeps
    the previous local PDF/DOCX fallback behavior for direct TextDocumentParser
    callers.
    """

    supported_extensions = SUPPORTED_ALL_EXTENSIONS

    def parse_bytes(
        self,
        content: bytes,
        *,
        filename: str,
        metadata: dict[str, Any] | None = None,
    ) -> ParsedDocument:
        suffix = Path(filename).suffix.lower()
        self._ensure_supported(suffix)

        if suffix in SUPPORTED_TEXT_EXTENSIONS:
            return super().parse_bytes(content, filename=filename, metadata=metadata)
        if suffix == ".pdf":
            text = MinerUDocumentParser._parse_pdf_local(content)
        elif suffix == ".docx":
            text = MinerUDocumentParser._parse_docx_local(content)
        else:
            supported = ", ".join(sorted(self.supported_extensions))
            raise ValueError(f"Unsupported document type: {suffix}. Supported: {supported}")

        return ParsedDocument(
            filename=filename,
            mime_type=MIME_TYPES.get(suffix, "text/plain"),
            text=self._normalize_newlines(text),
            metadata={**(metadata or {}), "extension": suffix, "parser": "local"},
        )
 
